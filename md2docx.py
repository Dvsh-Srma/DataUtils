import re
import argparse
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

# Global dictionary for font sizes
FONT_SIZES = {
    'heading1': Pt(28),
    'heading2': Pt(24),
    'heading3': Pt(20),
    'heading4': Pt(16),
    'heading5': Pt(14),
    'heading6': Pt(12),
    'normal': Pt(12),
    'code': Pt(10)
}

def add_hyperlink(paragraph, text, url):
    """
    A function that adds a hyperlink to a paragraph.
    This uses a low-level workaround to insert an external
    hyperlink into the document.
    """
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), "0000FF")
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), "single")
    rPr.append(underline)
    new_run.append(rPr)

    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def process_inline_formatting(paragraph, text):
    """
    Process inline markdown formatting within a text line and
    add runs to the paragraph.

    Supported inline elements:
      - Hyperlinks: [text](url)
      - Bold: **text**
      - Inline code: `code`
      - Italic: *text*
    """
    pattern = re.compile(
        r'(?P<link>\[(?P<link_text>.*?)\]\((?P<link_url>.*?)\))'
        r'|(?P<bold>\*\*(?P<bold_text>.*?)\*\*)'
        r'|(?P<code>`(?P<code_text>.*?)`)'
        r'|(?P<italic>\*(?P<italic_text>.*?)\*)'
    )
    pos = 0
    for match in pattern.finditer(text):
        start, end = match.span()
        if start > pos:
            paragraph.add_run(text[pos:start])
        if match.group('link'):
            add_hyperlink(paragraph, match.group('link_text'), match.group('link_url'))
        elif match.group('bold'):
            run = paragraph.add_run(match.group('bold_text'))
            run.bold = True
        elif match.group('code'):
            run = paragraph.add_run(match.group('code_text'))
            run.font.name = "Courier New"
            run.font.size = FONT_SIZES['code']
        elif match.group('italic'):
            run = paragraph.add_run(match.group('italic_text'))
            run.italic = True
        pos = end
    if pos < len(text):
        paragraph.add_run(text[pos:])

def set_paragraph_font_size(paragraph, font_size):
    """
    Set the font size for all runs in a paragraph that do not
    have a size defined.
    """
    for run in paragraph.runs:
        if run.font.size is None:
            run.font.size = font_size

def convert_markdown_to_docx(markdown_text, output_filename):
    """
    Convert markdown text to a well-formatted DOCX file.
    Supports:
      - Code blocks (wrapped in triple backticks)
      - Tables (using markdown table syntax)
      - Images (if a line shows only an image markdown)
      - Headings, bullet lists, and inline formatting (bold, italic,
        inline code, hyperlinks)
    """
    doc = Document()
    lines = markdown_text.splitlines()
    i = 0
    in_code_block = False
    code_block_lines = []

    while i < len(lines):
        line = lines[i]

        # Toggle code block mode based on ```
        if line.strip().startswith("```"):
            if in_code_block:
                # End code block: add collected code lines with monospace font and code size.
                p = doc.add_paragraph()
                for code_line in code_block_lines:
                    run = p.add_run(code_line + "\n")
                    run.font.name = "Courier New"
                    run.font.size = FONT_SIZES['code']
                in_code_block = False
                code_block_lines = []
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_block_lines.append(line)
            i += 1
            continue

        # Table detection: current line with "|" and next line as separator.
        if '|' in line and i + 1 < len(lines) and re.match(r'^\s*\|?(?:\s*:?-+:?\s*\|)+\s*(?:\|)?\s*$', lines[i+1]):
            header_line = line
            table_data = []
            header_cells = [cell.strip() for cell in header_line.strip().strip('|').split('|')]
            table_data.append(header_cells)
            i += 2  # Skip header and separator.
            while i < len(lines) and '|' in lines[i]:
                row_cells = [cell.strip() for cell in lines[i].strip().strip('|').split('|')]
                table_data.append(row_cells)
                i += 1
            if table_data:
                rows_count = len(table_data)
                cols_count = len(table_data[0])
                table = doc.add_table(rows=rows_count, cols=cols_count)
                table.style = 'Table Grid'
                for r, row in enumerate(table_data):
                    for c, cell in enumerate(row):
                        p_cell = table.rows[r].cells[c].paragraphs[0]
                        process_inline_formatting(p_cell, cell)
                        set_paragraph_font_size(p_cell, FONT_SIZES['normal'])
            continue

        # Image detection: if the entire line is an image.
        image_match = re.match(r'^\s*!\[(.*?)\]\((.*?)\)\s*$', line)
        if image_match:
            alt_text = image_match.group(1)
            image_path = image_match.group(2)
            try:
                doc.add_picture(image_path)
            except Exception as e:
                p = doc.add_paragraph()
                p.add_run(f"[Image '{alt_text}' could not be loaded: {str(e)}]")
            i += 1
            continue

        # Markdown headers.
        header_match = re.match(r'^(#{1,6})\s+(.*)', line)
        if header_match:
            level = len(header_match.group(1))
            text = header_match.group(2)
            # Create the heading paragraph and process inline formatting.
            heading_paragraph = doc.add_heading(level=level)
            process_inline_formatting(heading_paragraph, text)
            set_paragraph_font_size(heading_paragraph, FONT_SIZES.get(f'heading{level}', FONT_SIZES['normal']))
            i += 1
            continue

        # Unordered bullet list items.
        list_match = re.match(r'^[-*]\s+(.*)', line)
        if list_match:
            text = list_match.group(1)
            p = doc.add_paragraph(style='List Bullet')
            process_inline_formatting(p, text)
            set_paragraph_font_size(p, FONT_SIZES['normal'])
            i += 1
            continue

        # Blank line.
        if line.strip() == "":
            doc.add_paragraph("")
            i += 1
            continue

        # Default paragraph.
        p = doc.add_paragraph()
        process_inline_formatting(p, line)
        set_paragraph_font_size(p, FONT_SIZES['normal'])
        i += 1

    doc.save(output_filename)
    print(f"DOCX file saved as: {output_filename}")

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown text or a Markdown file to a formatted DOCX file."
    )
    parser.add_argument(
        "-i", "--input", type=str,
        help="Input Markdown file path. Defaults to 'markdownresp.txt' if not provided."
    )
    parser.add_argument(
        "-o", "--output", type=str,
        help="Output DOCX file path. Defaults to 'output.docx' if not provided."
    )
    args = parser.parse_args()

    try:
        with open(args.input, "r", encoding="utf-8") as f:
            markdown_text = f.read()
    except FileNotFoundError:
        print(f"Error: Input file '{args.input}' not found.")
        return

    convert_markdown_to_docx(markdown_text, args.output)

if __name__ == "__main__":
    main()